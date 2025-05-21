import streamlit as st
import subprocess
import whisper
import os
import re
import tempfile
from docx import Document

def baixar_audio(video_url: str, saida_audio: str):
    """Baixa o áudio como MP3 para compatibilidade no Streamlit Cloud"""
    result = subprocess.run([
        "yt-dlp",
        "--extract-audio",
        "--audio-format", "mp3",
        "-o", saida_audio,
        video_url
    ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

    if result.returncode != 0:
        raise Exception(f"Erro ao baixar áudio:\n{result.stderr}")

def transcrever_audio(modelo: str, arquivo_audio: str) -> str:
    model = whisper.load_model(modelo)
    result = model.transcribe(arquivo_audio, language="pt")
    return result["text"]

def salvar_txt(texto: str, caminho: str):
    with open(caminho, "w", encoding="utf-8") as f:
        f.write(texto)

def formatar_para_word(texto: str, caminho_docx: str):
    paragrafos = re.split(r'(?<=[.!?])\s*\n*', texto.strip())
    doc = Document()
    doc.add_heading("Transcrição Formatada", level=1)
    for p in paragrafos:
        if p.strip():
            doc.add_paragraph(p.strip())
            doc.add_paragraph("")
    doc.save(caminho_docx)

# ------------------ INTERFACE STREAMLIT ------------------ #

st.markdown("""
<h1 style='text-align: center;'>
🎧 Transcritor de Vídeo YouTube com Whisper
</h1>
""", unsafe_allow_html=True)

video_url = st.text_input("Cole o link do vídeo do YouTube:")

if st.button("Transcrever"):
    if not video_url:
        st.warning("Por favor, cole um link válido.")
    else:
        try:
            progress = st.progress(0)
            with tempfile.TemporaryDirectory() as temp_dir:
                audio_path = os.path.join(temp_dir, "audio.%(ext)s")  # yt-dlp substitui %(ext)s
                audio_mp3 = os.path.join(temp_dir, "audio.mp3")
                txt_path = os.path.join(temp_dir, "transcricao.txt")
                docx_path = os.path.join(temp_dir, "transcricao.docx")
                
                progress.progress(10)
                st.info("🔄 Baixando áudio...")
                baixar_audio(video_url, audio_path)

                # Após download, yt-dlp deve criar algo como audio.mp3
                downloaded_file = [f for f in os.listdir(temp_dir) if f.endswith(".mp3")][0]
                downloaded_audio_path = os.path.join(temp_dir, downloaded_file)
                
                progress.progress(40)
                st.info("🧠 Transcrevendo áudio com Whisper...")
                texto = transcrever_audio("base", downloaded_audio_path)
                
                progress.progress(70)
                st.info("💾 Salvando transcrição...")
                salvar_txt(texto, txt_path)
                
                progress.progress(90)
                st.info("📄 Formatando transcrição para .docx...")
                formatar_para_word(texto, docx_path)
                
                progress.progress(100)
                st.success("✅ Transcrição concluída!")

                with open(docx_path, "rb") as f:
                    st.download_button("📥 Baixar arquivo .docx", f, file_name="transcricao.docx")
                    progress.empty()

        except Exception as e:
            st.error(f"❌ Ocorreu um erro: {e}")

st.markdown("""
<hr style="margin-top: 3rem; margin-bottom: 1rem;">
<p style='text-align: center; font-size: 0.875rem; color: gray;'>
Desenvolvido por <b>Ronald César</b> • Versão 1.0
</p>
""", unsafe_allow_html=True)
