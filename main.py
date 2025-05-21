import subprocess
import whisper
import os
import re
from docx import Document
from datetime import datetime

def baixar_audio(video_url: str, saida_audio: str):
    print("🎥 Baixando o áudio do vídeo...")
    try:
        subprocess.run([
            "yt-dlp", "-f", "bestaudio",
            "--extract-audio", "--audio-format", "mp3",
            "-o", saida_audio, video_url
        ], check=True)
        print(f"✅ Áudio baixado como: {saida_audio}")
    except subprocess.CalledProcessError:
        print("❌ Erro ao baixar o vídeo. Verifique o link ou sua conexão.")
        exit(1)

def transcrever_audio(modelo: str, arquivo_audio: str) -> str:
    print("🧠 Carregando modelo Whisper e transcrevendo...")
    try:
        model = whisper.load_model(modelo)
        result = model.transcribe(arquivo_audio, language="pt")
        return result["text"]
    except Exception as e:
        print(f"❌ Erro durante transcrição: {e}")
        exit(1)

def salvar_txt(texto: str, caminho: str):
    with open(caminho, "w", encoding="utf-8") as f:
        f.write(texto)
    print(f"✅ Transcrição salva como: {caminho}")

def formatar_para_word(texto: str, caminho_docx: str):
    print("📝 Formatando para .docx...")
    # Divide com base em pontuação final (., !, ?) e quebra de linha
    paragrafos = re.split(r'(?<=[.!?])\s*\n*', texto.strip())
    
    doc = Document()
    doc.add_heading("Transcrição Formatada", level=1)

    for p in paragrafos:
        if p.strip():
            doc.add_paragraph(p.strip())
            doc.add_paragraph("")  # espaço entre parágrafos

    doc.save(caminho_docx)
    print(f"✅ Arquivo .docx criado: {caminho_docx}")

def main():
    # 1. Link do vídeo
    video_url = "https://www.youtube.com/watch?v=rFizDJnFW4w"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    audio_filename = f"audio_{timestamp}.mp3"
    txt_filename = f"transcricao_{timestamp}.txt"
    docx_filename = f"transcricao_{timestamp}.docx"

    # 2. Baixar áudio
    baixar_audio(video_url, audio_filename)

    # 3. Transcrever
    texto = transcrever_audio("base", audio_filename)

    # 4. Salvar .txt
    salvar_txt(texto, txt_filename)

    # 5. Salvar .docx
    formatar_para_word(texto, docx_filename)

if __name__ == "__main__":
    main()
